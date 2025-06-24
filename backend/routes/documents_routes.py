import os
import re
import hashlib
import difflib
from datetime import datetime
from flask import Blueprint, request, jsonify, send_from_directory,current_app,Response
from werkzeug.utils import secure_filename
from bson import ObjectId
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from extensions import mongo
from utils import token_required
from langchain_ollama import OllamaEmbeddings  # import embeddings
from utils.faiss_index import delete_from_index, update_faiss_index,CachedEmbeddingModel
from data_ingestion import (search_arxiv, search_pubmed,
                            search_openalex, merge_and_deduplicate)  # fonctions de recherche
from langchain_community.vectorstores import FAISS
import requests
import pdfplumber
import langid
import pandas as pd
import markdown
import PyPDF2
import subprocess
from utils.paths import get_user_folder, get_user_index_folder
from PIL import Image
from utils.extract_text_from_pdf import extract_pdf_content
import nltk
import docx2txt
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import pytesseract
from pdf2image import convert_from_path
from bs4 import BeautifulSoup
from .multi_functions_routes import ollama_query
from deep_translator import GoogleTranslator


bp_documents = Blueprint("documents", __name__)
BASE_DIR = os.path.dirname(__file__)
UPLOAD_FOLDER = "documents"
THUMBNAIL_FOLDER = "thumbnails"
ALLOWED_EXTENSIONS = {
    "pdf", "txt", "doc", "docx", "rtf", 
    "xls", "xlsx", "csv", "ppt", "pptx", 
    "png", "jpg", "jpeg", "tiff", 
    "xml", "json", "html", "md"
}
EXTERNAL_UPLOAD_FOLDER = "./external_uploads"



def translate_text(text, target_lang):
    try:
        return GoogleTranslator(source='auto', target=target_lang).translate(text)
    except Exception:
        return text
    
def find_best_fuzzy_match(haystack, needle, min_ratio=0.6):
    needle_len = len(needle)
    best_start = -1
    best_ratio = 0
    best_substring = None

    for start in range(len(haystack) - needle_len + 1):
        substring = haystack[start : start + needle_len]
        ratio = difflib.SequenceMatcher(None, needle, substring).ratio()
        if ratio > best_ratio and ratio >= min_ratio:
            best_ratio = ratio
            best_start = start
            best_substring = substring

    if best_start == -1:
        return None, None, None
    else:
        return best_substring, best_start, best_start + len(best_substring)
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
def compute_file_hash(filepath):
    """Compute SHA256 hash of a file to prevent duplicates"""
    hash_sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_sha256.update(chunk)
    return hash_sha256.hexdigest()

def clean_text(text):
    """Advanced text cleaning"""
    # Remove non-printable characters
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', ' ', text)
    
    # Remove common headers/footers
    patterns = [
        r'\b(page|confidential|draft)\b.*\n',
        r'\d{1,2}/\d{1,2}/\d{2,4}',
        r'©.*\n',
        r'http[s]?://\S+',
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    ]
    
    for pattern in patterns:
        text = re.sub(pattern, ' ', text, flags=re.IGNORECASE)
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
def create_thumbnail(filepath, save_path, size=(200, 200)):
    ext = filepath.rsplit('.', 1)[1].lower()
    try:
        if ext == 'pdf':
            # Extraire la première page du PDF comme image
            with pdfplumber.open(filepath) as pdf:
                page = pdf.pages[0]
                pil_image = page.to_image(resolution=150).original
                pil_image.thumbnail(size)
                pil_image.save(save_path)
        elif ext in ['jpg', 'jpeg', 'png', 'gif']:
            img = Image.open(filepath)
            img.thumbnail(size)
            img.save(save_path)
        elif ext in [".docx"]:
            # Pour DOCX, générer une thumbnail textuelle
            text = docx2txt.process(filepath)
            if text:
                img = Image.new("RGB", (400, 300), color=(240, 240, 240))
                from PIL import ImageDraw, ImageFont
                draw = ImageDraw.Draw(img)
                draw.text((10, 10), text[:300] + "...", fill=(0, 0, 0))
                img.save(save_path, "PNG")
        else:
            # Pas de thumbnail pour autres fichiers
            pass
    except Exception as e:
        print(f"Erreur création thumbnail pour {filepath} : {e}")

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = page_text.strip()
                    text += cleaned_text + "\n"
    except Exception as e:
        print(f"Erreur extraction PDF {filepath} : {e}")
    
    if not text.strip():
        try:
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text.strip() + "\n"
        except Exception as e2:
            print(f"PyPDF2 aussi a échoué sur {filepath} : {e2}")
    
    return text.strip()

def extract_text_with_ocr(filepath):
    """OCR extraction with automatic language detection"""
    try:

        
        lang_map = {
            'en': 'eng',
            'fr': 'fra',
            'es': 'spa',
            'de': 'deu',
            'it': 'ita'
        }
        
        text = ""
        if filepath.lower().endswith('.pdf'):
            images = convert_from_path(filepath)
        else:
            images = [Image.open(filepath)]
        
        for img in images:
            # Language detection on sample
            sample_text = pytesseract.image_to_string(img, lang='eng+fra')
            lang, _ = langid.classify(sample_text[:500] if sample_text else 'en')
            lang_code = lang_map.get(lang, 'eng')
            
            # OCR with detected language
            text += pytesseract.image_to_string(img, lang=lang_code) + "\n\n"
        
        return clean_text(text)
    except Exception as e:
        print(f"Erreur OCR: {e}")
        return ""

def extract_office_document(filepath):
    """Extraction robuste pour documents Word"""
    import subprocess
    from docx import Document

    ext = os.path.splitext(filepath)[1].lower()
    text = ""

    try:
        if ext == ".docx":
            # ✅ Extraction avec python-docx
            doc = Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            print(f"[DOCX] Texte extrait depuis {filepath} : {len(text)} caractères")

        elif ext == ".doc":
            # ✅ Extraction avec antiword
            try:
                result = subprocess.run(
                    ["antiword", filepath],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                text = result.stdout
                print(f"[DOC] Texte extrait via antiword ({len(text)} caractères)")
            except Exception as e:
                print(f"[DOC] Erreur antiword : {e}")
                text = ""

        elif ext == ".rtf":
            # Optionnel : support RTF
            try:
                import striprtf
                with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                    rtf_text = f.read()
                    text = striprtf.rtf_to_text(rtf_text)
            except Exception as e:
                print(f"[RTF] Erreur striprtf : {e}")
                text = ""

        if not text.strip():
            print(f"[⚠️] Aucun texte extrait → fallback OCR")
            return extract_text_with_ocr(filepath)

        return clean_text(text)

    except Exception as e:
        print(f"[Office] ⚠️ Erreur extraction Office : {e}")
        return extract_text_with_ocr(filepath)
def extract_spreadsheet(filepath):
    """Extraction for spreadsheets"""
    try:
        text = ""
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext in ['.xls', '.xlsx']:
            xl = pd.ExcelFile(filepath)
            for sheet_name in xl.sheet_names:
                df = xl.parse(sheet_name)
                text += f"\n\n--- Feuille: {sheet_name} ---\n"
                text += df.to_string(index=False, max_rows=20)
        
        elif ext == '.csv':
            df = pd.read_csv(filepath, nrows=100)  # Limit to 100 rows
            text = df.to_string(index=False)
        
        return clean_text(text)
    except Exception as e:
        print(f"Erreur extraction tableur: {e}")
        return ""

# def extract_presentation(filepath):
#     """Extraction for PowerPoint presentations"""
#     try:
#         from pptx import Presentation
#         text = ""
#         prs = Presentation(filepath)
        
#         for i, slide in enumerate(prs.slides):
#             text += f"\n\n--- Slide {i+1} ---\n"
#             for shape in slide.shapes:
#                 if hasattr(shape, "text"):
#                     text += shape.text + "\n"
        
#         return clean_text(text)
#     except Exception as e:
#         print(f"Erreur extraction présentation: {e}")
#         return extract_text_with_ocr(filepath)

def extract_structured_content(filepath, format_type):
    """Extraction for structured formats"""
    content = ""
    try:
        if format_type == "xml":
            import xml.etree.ElementTree as ET
            tree = ET.parse(filepath)
            root = tree.getroot()
            
            if root.tag.endswith('}article'):  # JATS
                for elem in root.findall('.//abstract'):
                    content += elem.text + "\n"
                for elem in root.findall('.//body'):
                    content += ET.tostring(elem, encoding='unicode', method='text')
            elif root.tag == "TEI":  # TEI XML
                for elem in root.findall('.//text'):
                    content += ET.tostring(elem, encoding='unicode', method='text')
            else:  # Generic XML
                content = ET.tostring(root, encoding='unicode', method='text')
        
        elif format_type == "json":
            import json
            with open(filepath, 'r') as f:
                data = json.load(f)
                
                def extract_values(obj):
                    if isinstance(obj, dict):
                        for v in obj.values():
                            yield from extract_values(v)
                    elif isinstance(obj, list):
                        for item in obj:
                            yield from extract_values(item)
                    else:
                        yield str(obj)
                
                content = " ".join(extract_values(data))
        
        elif format_type == "html":
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f, 'html.parser')
                content = soup.get_text(separator='\n', strip=True)
        
        elif format_type == "md":
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                md_text = f.read()
                html = markdown.markdown(md_text)
                soup = BeautifulSoup(html, 'html.parser')
                content = soup.get_text(separator='\n', strip=True)
    
    except Exception as e:
        print(f"Erreur extraction structurée: {e}")
    
    return content
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os

def extract_file_content(filepath):
    """Extraction générique du contenu d'un fichier, avec fallback OCR si besoin"""
    ext = os.path.splitext(filepath)[1].lower().replace('.', '')
    content = ""

    try:
        if ext == 'pdf':
            content = extract_text_from_pdf(filepath)
            if not content.strip() or len(content) < 100:
                print("[INFO] Faible contenu PDF – passage en OCR")
                content = extract_text_with_ocr(filepath)

        elif ext in ['png', 'jpg', 'jpeg', 'tiff']:
            content = extract_text_with_ocr(filepath)

        elif ext in ['doc', 'docx', 'rtf']:
            content = extract_office_document(filepath)

        elif ext in ['xls', 'xlsx', 'csv']:
            content = extract_spreadsheet(filepath)

        elif ext in ['xml', 'json', 'html', 'md']:
            content = extract_structured_content(filepath, ext)

        else:  # txt ou inconnu
            try:
                with open(filepath, "r", encoding="utf-8", errors="replace") as f:
                    content = f.read()
            except Exception as e:
                print(f"[ERREUR] Lecture fichier brut : {e}")
                content = extract_text_with_ocr(filepath)

    except Exception as e:
        print(f"[ERREUR] Extraction échouée ({ext}) : {e}")
        content = extract_text_with_ocr(filepath)

    if not content.strip():
        print("[⚠️] Aucun contenu extrait après fallback.")
        return []

    # ✅ Nettoyage et découpage
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
        add_start_index=True,
    )
    chunks = text_splitter.split_text(clean_text(content))
    print(f"[✅] {len(chunks)} chunks générés à partir du fichier : {os.path.basename(filepath)}")

    return chunks


@bp_documents.route("/upload", methods=["POST"])
@token_required
def upload_documents(user):
    data = request.form or {}
    conversation_id = data.get("conversation_id")
    if not conversation_id:
        return jsonify({"error": "conversation_id manquant"}), 400

    if "files" not in request.files:
        return jsonify({"error": "Aucun fichier"}), 400

    files = request.files.getlist("files")
    saved_files = []

    user_folder = os.path.join(UPLOAD_FOLDER, str(user["_id"]))
    os.makedirs(user_folder, exist_ok=True)

    new_documents = []

    for f in files:
        if f and allowed_file(f.filename):
            filename = secure_filename(f.filename)
            filepath = os.path.join(user_folder, filename)
            f.save(filepath)
            chunks = extract_file_content(filepath)

            if chunks:
                user_thumb_folder = os.path.join(THUMBNAIL_FOLDER, str(user["_id"]))
                os.makedirs(user_thumb_folder, exist_ok=True)
                thumb_filename = f"{os.path.splitext(filename)[0]}.png"
                thumb_path = os.path.join(user_thumb_folder, thumb_filename)
                create_thumbnail(filepath, thumb_path)

                file_id = ObjectId()
                for i, chunk in enumerate(chunks):
                    doc_entry = {
                        "filename": filename,
                        "path": filepath,
                        "user_id": user["_id"],
                        "conversation_id": conversation_id,
                        "content": chunk,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "file_type": os.path.splitext(filename)[1][1:].upper(),
                        "upload_date": datetime.utcnow(),
                        "thumbnail": thumb_filename,
                        "is_master": i == 0,
                        "file_id": file_id,
                    }
                    inserted = mongo.db.documents.insert_one(doc_entry)
                    if i == 0:
                        saved_files.append(doc_entry)

                    new_documents.append({
                        "id": str(inserted.inserted_id),
                        "content": chunk,
                        "metadata": {
                            "filename": filename,
                            "chunk_index": i,
                            "conversation_id": conversation_id,
                            "file_type": doc_entry["file_type"],
                            "file_id": str(file_id),
                            "is_master": i == 0
                        }
                    })

    if new_documents:
        index_folder = os.path.join(user_folder, "conversations", str(conversation_id))
        os.makedirs(index_folder, exist_ok=True)
        update_faiss_index(index_folder, new_documents, user_id=user["_id"])

    return jsonify({"files": saved_files})


import os

@bp_documents.route("/download_external_document", methods=["POST"])
@token_required
def download_external_document(user):
    def arxiv_to_pdf_url(url):
        if "arxiv.org/abs/" in url:
            return url.replace("arxiv.org/abs/", "arxiv.org/pdf/") + ".pdf"
        return url

    data = request.get_json()
    file_url = data.get("file_url")
    filename = data.get("filename")
    source = data.get("source", "unknown")

    if not file_url or not filename:
        return jsonify({"error": "URL ou nom de fichier manquant"}), 400

    file_url = arxiv_to_pdf_url(file_url)

    if not filename.lower().endswith(".pdf"):
        filename += ".pdf"

    try:
        response = requests.get(file_url, allow_redirects=True, timeout=10)
        if response.status_code != 200:
            return jsonify({"error": f"Erreur téléchargement : {response.status_code}"}), 400

        content_type = response.headers.get("Content-Type", "")
        if not content_type.startswith("application/pdf"):
            return jsonify({"error": "Le fichier téléchargé n'est pas un PDF valide."}), 400

        user_folder = os.path.join(EXTERNAL_UPLOAD_FOLDER, str(user["_id"]))
        os.makedirs(user_folder, exist_ok=True)
        filepath = os.path.join(user_folder, filename)

        with open(filepath, "wb") as f:
            f.write(response.content)
        print(f"[Téléchargement] Fichier enregistré à : {filepath}")

        content_data = extract_pdf_content(filepath)
        text_content = content_data["text_sections"][:100_000]

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=50)
        chunks = text_splitter.split_text(text_content)[:200]
        print(f"[Chunking] Nombre de chunks générés : {len(chunks)}")

        documents = [
            Document(page_content=chunk, metadata={"source": f"{filename}_chunk_{i}"})
            for i, chunk in enumerate(chunks)
        ]
        index_path = os.path.join(user_folder, "faiss_index")

        embedding = CachedEmbeddingModel(user["_id"])

        try:
            store = FAISS.load_local(index_path, embeddings=embedding, allow_dangerous_deserialization=True)
            store.add_documents(documents)
            print("[FAISS] Index existant mis à jour.")
        except Exception:
            store = FAISS.from_documents(documents, embedding=embedding)
            print("[FAISS] Nouvel index créé.")

        store.save_local(index_path)
        print("[FAISS] Index sauvegardé.")

        original_url = data.get("original_url", file_url)

        doc_entry = {
            "_id": original_url,
            "title": filename,
            "path": filepath,
            "source": source,
            "content": text_content,
            "tables_count": content_data["tables_count"],
            "figures_count": content_data["figures_count"],
            "upload_date": datetime.utcnow(),
            "indexed": True,
            "url": file_url,
        }

        mongo.db.documents_externes.replace_one(
            {"_id": original_url},
            doc_entry,
            upsert=True
        )

        inserted_doc = mongo.db.documents_externes.find_one({"_id": original_url})
        print("[MongoDB] Document inséré ou mis à jour :", inserted_doc)

        return jsonify({"message": "Téléchargement et indexation OK", "document": doc_entry})

    except Exception as e:
        print(f"[Erreur] Exception lors de l’indexation : {e}")
        return jsonify({"error": f"Erreur serveur : {str(e)}"}), 500



@bp_documents.route("/document_chunks", methods=["POST"])
@token_required
def get_document_chunks(user):
    try:
        data = request.get_json()
        document_id = data.get("document_id")
        language = data.get("language", "en")  # Récupère la langue, défaut 'en'

        if not document_id:
            return jsonify({"error": "document_id requis"}), 400

        # Conversion URL arXiv
        if "arxiv.org/abs/" in document_id:
            document_id = document_id.replace("arxiv.org/abs/", "arxiv.org/pdf/") + ".pdf"

        # Recherche en MongoDB
        doc = mongo.db.documents_externes.find_one({"_id": document_id})

        # Téléchargement + extraction si absent
        if not doc:
            filepath = os.path.join(EXTERNAL_UPLOAD_FOLDER, str(user["_id"]), os.path.basename(document_id))
            os.makedirs(os.path.dirname(filepath), exist_ok=True)

            if not os.path.exists(filepath):
                response = requests.get(document_id, stream=True, timeout=15)
                response.raise_for_status()

                content_type = response.headers.get("Content-Type", "")
                if not content_type.startswith("application/pdf"):
                    return jsonify({"error": "Le fichier distant n'est pas un PDF valide."}), 400

                with open(filepath, "wb") as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)

            content_data = extract_pdf_content(filepath)
            text_content = content_data["text_sections"][:100_000]

            doc = {
                "_id": document_id,
                "title": os.path.basename(filepath),
                "path": filepath,
                "source": "unknown",
                "content": text_content,
                "tables_count": content_data["tables_count"],
                "figures_count": content_data["figures_count"],
                "upload_date": datetime.utcnow(),
                "indexed": False,
                "url": document_id
            }
            mongo.db.documents_externes.insert_one(doc)

        # Traduction du contenu si nécessaire
        content_to_use = doc.get("content", "")
        if language != "en":
            try:
                content_to_use = translate_text(content_to_use, language)
            except Exception as e:
                print(f"[⚠️ Traduction contenu échouée] : {e}")

        # Indexation FAISS (sans traduction, on indexe le texte original)
        if not doc.get("indexed", False):
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=50)
            chunks = text_splitter.split_text(doc["content"])[:200]

            documents = [
                Document(page_content=chunk, metadata={"source": f"{doc['title']}_chunk_{i}"})
                for i, chunk in enumerate(chunks)
            ]

            user_folder = os.path.join(EXTERNAL_UPLOAD_FOLDER, str(user["_id"]))
            os.makedirs(user_folder, exist_ok=True)
            index_path = os.path.join(user_folder, "faiss_index")

            embedding = CachedEmbeddingModel(user["_id"])

            try:
                store = FAISS.load_local(index_path, embeddings=embedding, allow_dangerous_deserialization=True)
                store.add_documents(documents)
            except Exception:
                store = FAISS.from_documents(documents, embedding=embedding)

            store.save_local(index_path)

            mongo.db.documents_externes.update_one(
                {"_id": document_id},
                {"$set": {"indexed": True}}
            )

        # Si on a déjà highlight_chunks en base, renvoyer
        if "highlight_chunks" in doc:
            full_text = " ".join(doc.get("content", "").splitlines())[:20000]
            valid_chunks = [chunk for chunk in doc["highlight_chunks"] if chunk.get("text") in full_text]

            # Traduction des extraits avant renvoi
            if language != "en":
                for chunk in valid_chunks:
                    try:
                        chunk["text"] = translate_text(chunk["text"], language)
                        if chunk.get("label"):
                            chunk["label"] = translate_text(chunk["label"], language).upper()
                    except Exception as e:
                        print(f"[⚠️ Traduction highlight_chunks échouée] : {e}")

            return jsonify({
                "chunks": valid_chunks,
                "full_text": content_to_use[:20000]  # texte potentiellement traduit
            })

        # Construction prompt avec contenu traduit (ou original si en='en')
        prompt = f"""
        Tu es un expert en lecture de documents scientifiques. Voici le contenu partiel d’un document.
        Identifie les passages les plus pertinents concernant l’objectif, les méthodes, les résultats et la conclusion.

        Pour chaque passage trouvé, ajoute un titre en majuscules (ex: OBJECTIF, METHODE, RESULTATS, CONCLUSION).
        Donne-moi une liste de 3 à 5 extraits pertinents dans ce format :

        OBJECTIF:
        [passage]

        METHODE:
        [passage]

        ...

        Document :
        {content_to_use[:5000]}
        """

        response = ollama_query(prompt)
        raw_chunks = [chunk.strip() for chunk in response.strip().split("\n\n") if chunk.strip()]
        if raw_chunks and raw_chunks[0].lower().startswith("here are"):
            raw_chunks = raw_chunks[1:]

        full_text = " ".join(doc.get("content", "").splitlines())[:20000]

        chunks = [chunk.replace('\n', ' ').strip() for chunk in raw_chunks]

        valid_chunks_with_pos = []
        for chunk in chunks:
            label_match = re.match(r"^(OBJECTIF|METHODE|RESULTATS|CONCLUSION):", chunk, re.IGNORECASE)
            label = label_match.group(1).capitalize() if label_match else None
            clean_chunk = re.sub(r"^(OBJECTIF|METHODE|RESULTATS|CONCLUSION):", "", chunk, flags=re.IGNORECASE).strip()

            matched_text, start, end = find_best_fuzzy_match(full_text, clean_chunk, min_ratio=0.5)
            if matched_text:
                valid_chunks_with_pos.append({
                    "text": matched_text,
                    "start": start,
                    "end": end,
                    "label": label
                })
            else:
                print("No match found.")

        # Traduction des extraits trouvés avant sauvegarde et envoi
        if language != "en":
            for chunk in valid_chunks_with_pos:
                try:
                    chunk["text"] = translate_text(chunk["text"], language)
                    if chunk.get("label"):
                        chunk["label"] = translate_text(chunk["label"], language).upper()
                except Exception as e:
                    print(f"[⚠️ Traduction des extraits échouée] : {e}")

        mongo.db.documents_externes.update_one(
            {"_id": document_id},
            {"$set": {"highlight_chunks": valid_chunks_with_pos}}
        )

        return jsonify({
            "chunks": valid_chunks_with_pos,
            "full_text": content_to_use[:20000]
        })

    except Exception as e:
        print(f"[Erreur] Exception dans /document_chunks : {e}")
        return jsonify({"error": str(e)}), 500



@bp_documents.route('/api/favorites', methods=['POST'])
@token_required
def add_favorite(user):
    data = request.get_json()
    document_id = data.get('documentId')
    document_data = data.get('documentData')  # optionnel, pour infos du doc

    if not document_id:
        return jsonify({'error': 'documentId manquant'}), 400

    # Recherche document dans la base
    doc = mongo.db.documents_externes.find_one({"_id": document_id})

    if not doc:
        # Si document non trouvé, on l'ajoute à la base (avec infos minimales)
        # Si tu n'as pas toutes les données, prends au moins ce que tu peux
        if not document_data:
            return jsonify({'error': 'Document non trouvé et pas assez d\'infos pour l\'ajouter'}), 400

        # Conversion _id string en ObjectId
        document_data["_id"] = document_id
        mongo.db.documents_externes.insert_one(document_data)

    # Ajout aux favoris de l'utilisateur
    user_id = user["_id"]

    mongo.db.users_profile.update_one(
        {"_id": user_id},
        {"$addToSet": {"favorites": document_id}},
        upsert=True
    )

    return jsonify({'message': 'Document ajouté aux favoris'})


@bp_documents.route('/favorites/<document_id>', methods=['DELETE'])
@token_required
def remove_favorite(user, document_id):
    user_id = user["_id"]
    mongo.db.users.update_one(
        {"_id": user_id},
        {"$pull": {"favorites": document_id}}
    )
    return jsonify({'message': 'Document retiré des favoris'})


@bp_documents.route('/favorites', methods=['GET'])
@token_required
def get_favorites(user):
    user_id = user["_id"]

    # Récupérer la liste des _id favoris de l'utilisateur
    user_data = mongo.db.users_profile.find_one({"_id": user_id}, {"favorites": 1})
    favorites_ids = user_data.get("favorites", [])

    # Chercher les documents dans documents_externes
    favorites_docs = list(
        mongo.db.documents_externes.find({"_id": {"$in": favorites_ids}})
    )

    # Pour renvoyer au frontend, on peut transformer ObjectId en string si besoin
    for doc in favorites_docs:
        doc["_id"] = str(doc["_id"])

    return jsonify({"favorites": favorites_docs})


@bp_documents.route('/thumbnails/<user_id>/<filename>')
def serve_thumbnail(user_id, filename):
    thumb_path = os.path.join(THUMBNAIL_FOLDER, user_id)
    return send_from_directory(thumb_path, filename)


@bp_documents.route('/user_documents/<conversation_id>', methods=['GET'])
@token_required
def get_user_documents(user, conversation_id):
    docs = list(mongo.db.documents.find({"conversation_id": conversation_id, "user_id": user["_id"]}))
    for doc in docs:
        doc["_id"] = str(doc["_id"])
        if isinstance(doc["user_id"], ObjectId):
            doc["user_id"] = str(doc["user_id"])
    return jsonify({"files": docs})

@bp_documents.route("/documents", methods=["GET"])
@token_required
def get_user_all_documents(user):
    docs = list(mongo.db.documents.find({"user_id": user["_id"]}))
    
    formatted = []
    for doc in docs:
        upload_date = doc.get("upload_date")
        # Convertir en ISO string si upload_date existe
        upload_date_str = upload_date.isoformat() if upload_date else None
        
        formatted.append({
            "id": str(doc["_id"]),
            "title": doc.get("title", ""),
            "filename": doc.get("filename", ""),
            "uploaded_at": upload_date_str
        })

    return jsonify(formatted)


@bp_documents.route("/delete_document/<doc_id>", methods=["DELETE"])
@token_required
def delete_document(user, doc_id):
    try:
        object_id = ObjectId(doc_id)
    except Exception:
        return jsonify({"error": "ID de document invalide"}), 400

    doc = mongo.db.documents.find_one({"_id": object_id, "user_id": user["_id"]})
    if not doc:
        return jsonify({"error": "Document introuvable"}), 404

    file_id = doc.get("file_id")

    if file_id:
        # Récupérer tous les chunks liés
        chunks = list(mongo.db.documents.find({
            "file_id": file_id,
            "user_id": user["_id"]
        }))
        chunk_ids = [str(chunk["_id"]) for chunk in chunks]

        # Supprimer les documents (chunks) dans la base
        result = mongo.db.documents.delete_many({
            "file_id": file_id,
            "user_id": user["_id"]
        })

        # Supprimer le cache des embeddings associés
        try:
            mongo.db.embeddings_cache.delete_many({
                "chunk_id": {"$in": chunk_ids},
                "user_id": user["_id"]
            })
        except Exception as e:
            print(f"Erreur suppression cache embeddings : {e}")

        # Supprimer fichiers physiques et thumbnail si master
        if doc.get("is_master"):
            try:
                if os.path.exists(doc["path"]):
                    os.remove(doc["path"])
            except Exception as e:
                print(f"Erreur suppression fichier : {e}")

            try:
                thumb_path = os.path.join(
                    current_app.config["THUMBNAIL_FOLDER"],
                    str(user["_id"]),
                    doc["thumbnail"]
                )
                if os.path.exists(thumb_path):
                    os.remove(thumb_path)
            except Exception as e:
                print(f"Erreur suppression thumbnail : {e}")

            # Supprimer de l'index FAISS
            try:
                delete_from_index(user_id=user["_id"], doc_ids=chunk_ids)
            except Exception as e:
                print(f"Erreur suppression FAISS : {e}")

        return jsonify({
            "message": "Document, chunks, indexation et cache supprimés",
            "deleted_count": result.deleted_count
        }), 200

    # Fallback pour les documents anciens sans file_id
    mongo.db.documents.delete_one({"_id": object_id})
    mongo.db.embeddings_cache.delete_many({"chunk_id": str(object_id), "user_id": user["_id"]})

    return jsonify({"message": "Document supprimé (ancien format + cache)"}), 200



@bp_documents.route('/fetch_documents', methods=['POST'])
@token_required
def fetch_documents(user):
    data = request.get_json()
    query = data.get('query')
    source = data.get('source')
    max_results = data.get('max_results', 5)

    documents = []
    if source == 'all':
        all_docs = []
        all_docs.extend(search_arxiv(query, max_results))
        all_docs.extend(search_pubmed(query, max_results))
        # all_docs.extend(search_semantic_scholar(query, max_results))
        all_docs.extend(search_openalex(query, max_results))
        documents = merge_and_deduplicate(all_docs)
    elif source == 'arxiv':
        documents = search_arxiv(query, max_results)
    elif source == 'pubmed':
        documents = search_pubmed(query, max_results)
    # elif source == 'semantic_scholar':
    #     documents = search_semantic_scholar(query, max_results)
    elif source == 'openalex':
        documents = search_openalex(query, max_results)
    else:
        # gérer source inconnue si besoin
        documents = []

    documents = sorted(documents, key=lambda d: d.get('title', '').lower())
    return jsonify({'documents': documents})


@bp_documents.route("/proxy_pdf")
def proxy_pdf():
    url = request.args.get("url")
    if not url:
        return "Missing url parameter", 400
    try:
        r = requests.get(url, timeout=15)
        if r.status_code != 200 or "pdf" not in r.headers.get("Content-Type", ""):
            return "Not a valid PDF file", 400
        return Response(r.content, mimetype="application/pdf")
    except Exception as e:
        return str(e), 500